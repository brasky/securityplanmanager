from collections import defaultdict
from docx import *
from .models import Control, Implementation, Team, ControlOrigination, Certification
from django.db.models import Q
from .helper import get_control_parts
import sys
from django.core.exceptions import ObjectDoesNotExist

def get_customer_responsibility(implementation_details):
    teams = Team.objects.all()
    split_details = implementation_details.split('\n')
    # print(split_details)
    cust_resp_flag = False
    customer_responsibility = ''
    for part in split_details:
        if "Customer Responsibility:" in part:
            cust_resp_flag = True
            customer_responsibility = part
            continue
        if cust_resp_flag:
            for team in teams:
                if team.name + ":" in part or "Part 1" in part or "Part 2" in part or 'Part 1,2,3' in part:
                    return customer_responsibility
            customer_responsibility = customer_responsibility +'\n' + part

    return False

def get_part_text(implementation_details, control_parts):
    """
    get_part_text(string, list) -> string

    Takes implementation details string and control parts list and returns the text for the corresponding control part.
    """
    split_details = implementation_details.split('Part')
    part_text = ''
    part_num_colon = control_parts['part_num'] + ':'
    part_num_comma = control_parts['part_num'] + ','
    for part in split_details:
        if part_num_colon in part:
            part_text = part.lstrip('1234567890:, \n')
            return part_text
        elif part_num_comma in part:
            part_text = part.lstrip('1234567890:, \n') 
            return part_text


def create_implementation(new_implementation, certification):
    """
    create_implementation(dict) -> None

    Takes a dictionary containing Implementation fields, creates a new Implementation object, and adds it to a certification. Does not return anything. 
    """
    if new_implementation['customer_resp']:
        new_implementation_object = Implementation(
            control=new_implementation['control_object'],
            parameter=new_implementation['parameter'],
            responsible_role=new_implementation['responsible_role'],
            customer_responsibility=new_implementation['customer_resp'],
            solution=new_implementation['solution'],
            implementation_status=new_implementation['implementation_status'],
            
            )
    else:
        new_implementation_object = Implementation(
            control=new_implementation['control_object'],
            parameter=new_implementation['parameter'],
            responsible_role=new_implementation['responsible_role'],
            customer_responsibility='',
            solution=new_implementation['solution'],
            implementation_status=new_implementation['implementation_status'],
            
            )
    try:
        new_implementation_object.save()
        new_implementation_object.control_origination.set(new_implementation['control_origination'])
        new_implementation_object.teams.set(new_implementation['teams'])
        cert = Certification.objects.get(name__contains=certification)
        cert.implementations.add(new_implementation_object)
    except:
        e = sys.exc_info()[0]
        print( "Error: %s" % e )
        

def parse_solution_table(table, control_object, control_parts):
    """
    parse_solution_table(table object, control object, control parts list) -> implementation details string, customer responsibility string

    Takes the entire solution section of a control and splits the customer responsibility from the implementation details.
    """
    control_number = control_object.number
    try:
        if  not control_parts['part_letter'] and not control_parts['part_num']: #no letter, no part, no enhancement
            implementation_details = table.cell(0,1).text
            customer_responsibility = get_customer_responsibility(implementation_details)
            if customer_responsibility:
                implementation_details = implementation_details.replace(customer_responsibility, '').strip()

        elif not control_parts['part_num']:#letter, no part, no enhancement
            try:
                implementation_details = table.cell(control_parts['part_letter'],1).text
                customer_responsibility = get_customer_responsibility(implementation_details)
                if customer_responsibility:
                    implementation_details = implementation_details.replace(customer_responsibility, '').strip()
            except:
                print("parse solution table error:")
                # print(table.cell(4,0).text)
                print(control_number)
                print(control_parts)
                print(table.cell(0,0).text)
                return '', '' #blank implementation and customer responsibility

        else:#letter, part, no enhancement
            implementation_details = table.cell(control_parts['part_letter'],1).text
            customer_responsibility = get_customer_responsibility(implementation_details)
            if customer_responsibility:
                implementation_details = implementation_details.replace(customer_responsibility, '').strip()
        # else:
        #     raise ValueError('A case was not handled by the implementation table section of parse ssp')
    except IndexError:
        print("Table not correct in SSP for control: " + control_number)
        return "", ""  #there's a problem with how the table is laid out in the SSP. 
    return implementation_details, customer_responsibility

def get_implementation_status_from_cell(implementation_cell):
    """
    get_implementation_status_from_cell(Cell Object) -> String
    Determins which boxes are checked and returns the corresponding implementation status string.
    """
    implementation_status = ''
    for paragraph in implementation_cell.paragraphs:
        p = paragraph._element
        checkBoxes = p.xpath('.//w:checkBox')
        
        if 'w14:checked w14:val="1"' in p.xml:
            xpath_elements = p.xpath('.//w:t')
            implementation_status = xpath_elements[len(xpath_elements)-1].text.strip()
            if 'Partially' in implementation_status:
                implementation_status = 'PI'
            elif 'Implemented' in implementation_status:
                implementation_status = 'IM'
            elif 'Planned' in implementation_status:
                implementation_status = 'PL'
            elif 'Alternative' in implementation_status:
                implementation_status = 'AI'
            elif 'Not' in implementation_status:
                implementation_status = 'NA'
    return implementation_status


def get_control_origination_from_cell(control_origination_cell):
    """
    get_control_origination_from_cell(Cell object) -> list of control origination objects
    Determines which boxes are checked and returns an array of the corresponding control origination objects.
    """
    control_originations = []
    for paragraph in control_origination_cell.paragraphs:
        p = paragraph._element
        if 'w14:checked w14:val="1"' in p.xml:
            xpath_elements = p.xpath('.//w:t')
            control_origination = xpath_elements[len(xpath_elements)-1].text.strip()
            if "Service Provider Corporate" in control_origination:
                control_originations.append(ControlOrigination.objects.get(source='SPC'))
            elif "Service Provider System Specific" in control_origination:
                control_originations.append(ControlOrigination.objects.get(source='SPS'))
            elif "Hybrid" in control_origination:
                control_originations.append(ControlOrigination.objects.get(source='SPH'))
            elif "Configured" in control_origination:
                control_originations.append(ControlOrigination.objects.get(source='CBC'))
            elif "Provided" in control_origination:
                control_originations.append(ControlOrigination.objects.get(source='PBC'))
            elif "Shared" in control_origination:
                control_originations.append(ControlOrigination.objects.get(source='SHA'))
            elif "Inherited" in control_origination:
                control_originations.append(ControlOrigination.objects.get(source='INH'))
            elif "Not" in control_origination:
                control_originations.append(ControlOrigination.objects.get(source='NOT'))
    return control_originations


namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
implementations = []

def parse_ssp(file, certification):
    document = Document(file)
    for table in document.tables:
        try:
            table_title = table.cell(0,0).text
        except:
            continue
        if "Req" in  table_title or "req" in table_title:
            continue
            #currently Req. controls are not supported.
        if len(table_title) < 20 and '-' in table_title:
            control_parent = table_title
            parameters = defaultdict(str)
            controls = {}
            parameter = ''
            rows = len(table.rows)
            for cell in range(1,rows):

                if "Implementation" in table.cell(cell,0).text:
                    implementation_status = get_implementation_status_from_cell(table.cell(cell,0))

                elif "Control Origination" in table.cell(cell,0).text:
                    control_originations = get_control_origination_from_cell(table.cell(cell, 0))                  

                elif "Responsible Role" in table.cell(cell,0).text:
                    responsible_role = table.cell(cell,0).text.split(':')[1].strip()

                elif "Parameter" in table.cell(cell,0).text:
                    control = table.cell(cell, 0).text.split(':')[0]
                    parameter = table.cell(cell, 0).text.replace(control, '').strip(':').strip() #remove all the fluff, just the parameter will be stored.
                    control = control.replace('Parameter ', '').strip().replace('-0', '-').replace(' ', '')
                    if control.count('-') == 2: #there's more than one parameter for the control. i.e. Parameter AC-2 (2)-1: and Parameter AC-2 (2)-2: so we need to strip the -# from the end and concatenate the 2 together.
                        param_number = control[-1:] + ': '
                        control = control[:-2]
                        parameters[control] += param_number + parameter + '\n'
                    elif "-1(a)" in control:
                        parameters[control+"(1)"] += parameter + '\n'
                        parameters[control+"(2)"] += parameter + '\n'
                    else:
                        parameters[control] += parameter + '\n'
        
        elif "solution" in table_title:
            
            if len(control_parent) == 5: #if control in table title is a double digit control like AC-11
                control = control_parent.replace('-0', '-') + "(" #see if there are parts in the controls db like AC-11(a)
                matching_controls = Control.objects.filter(Q(number__contains=control) & ~Q(number__contains=' '))
                if not matching_controls: #if the query came back with no matching controls, then it must not have control parts in the solution table since there are no controls in the db.
                    control = control_parent.replace('-0', '-') #fix the formatting of the control to match the db.
                    matching_controls = Control.objects.filter(number=control)
                for control_object in matching_controls: #loop through the QuerySet that match. For example AC-11 would be AC-11(a) and AC-11(b) 
                    control_parts = get_control_parts(control_object.number)
                    implementation_details, customer_responsibility = parse_solution_table(table, control_object, control_parts)
                    if control_parts['part_num']:
                        implementation_details = get_part_text(implementation_details, control_parts)
                    parameter = parameters[control_object.number]
                    if implementation_details:
                        implementations_grouped_by_team = split_implementations(implementation_details)
                        for split_implementation in implementations_grouped_by_team:
                            teams, solution = split_implementation
                            new_implementation = {
                                'control_object': control_object,
                                'solution': solution,
                                'customer_resp': customer_responsibility,
                                'teams': teams,
                                'control_origination': control_originations,
                                'implementation_status': implementation_status,
                                'parameter': parameter,
                                'responsible_role': responsible_role
                            }
                            create_implementation(new_implementation, certification)

            elif len(control_parent) == 4:
                control = control_parent.replace('-0', '-') + "("
                matching_controls = Control.objects.filter(Q(number__contains=control) & ~Q(number__contains=' '))
                if not matching_controls:
                    control = control_parent.replace('-0', '-')
                    matching_controls = Control.objects.filter(number=control)
                for control_object in matching_controls:
                    control_parts = get_control_parts(control_object.number)
                    implementation_details, customer_responsibility = parse_solution_table(table, control_object, control_parts)
                    if control_parts['part_num']:
                        implementation_details = get_part_text(implementation_details, control_parts)
                    parameter = parameters[control_object.number]
                    if implementation_details:
                        implementations_grouped_by_team = split_implementations(implementation_details)
                        for split_implementation in implementations_grouped_by_team:
                            teams, solution = split_implementation
                            new_implementation = {
                                'control_object': control_object,
                                'solution': solution,
                                'customer_resp': customer_responsibility,
                                'teams': teams,
                                'control_origination': control_originations,
                                'implementation_status': implementation_status,
                                'parameter': parameter,
                                'responsible_role': responsible_role
                            }
                            create_implementation(new_implementation, certification)

            elif ' ' in control_parent and "Req." not in control_parent: #enhancements with spaces
                control = control_parent.replace('-0', '-')
                matching_controls = Control.objects.filter(number__contains=control)
                for control_object in matching_controls:
                    control_parts = get_control_parts(control_object.number)
                    implementation_details, customer_responsibility = parse_solution_table(table, control_object, control_parts)           
                    if control_parts['part_num']:
                        implementation_details = get_part_text(implementation_details, control_parts)
                    parameter = parameters[control_object.number.replace(' ', '')]
                    if implementation_details:
                        implementations_grouped_by_team = split_implementations(implementation_details)
                        for split_implementation in implementations_grouped_by_team:
                            teams, solution = split_implementation
                            new_implementation = {
                                'control_object': control_object,
                                'solution': solution,
                                'customer_resp': customer_responsibility,
                                'teams': teams,
                                'control_origination': control_originations,
                                'implementation_status': implementation_status,
                                'parameter': parameter,
                                'responsible_role': responsible_role
                            }
                            create_implementation(new_implementation, certification)
                    
            elif "Req." not in control_parent: #enhancements without spaces
                control_base = control_parent[:5]

                control_enhancement = control_parent[5:]

                control = control_base + ' ' + control_enhancement
                control = control.replace('-0', '-')

                matching_controls = Control.objects.filter(number__contains=control)
                if not matching_controls:
                    control_base = control_parent[:4]
                    control_enhancement = control_parent[4:]
                    control = control_base + ' ' + control_enhancement
                    control = control.replace('-0', '-')
                    matching_controls = Control.objects.filter(number__contains=control)
                for control_object in matching_controls:
                    control_parts = get_control_parts(control_object.number)
                    implementation_details, customer_responsibility = parse_solution_table(table, control_object, control_parts)           
                    if control_parts['part_num']:
                        implementation_details = get_part_text(implementation_details, control_parts)
                    parameter = parameters[control_object.number]
                    if implementation_details:
                        implementations_grouped_by_team = split_implementations(implementation_details)
                        for split_implementation in implementations_grouped_by_team:
                            teams, solution = split_implementation
                            new_implementation = {
                                'control_object': control_object,
                                'solution': solution,
                                'customer_resp': customer_responsibility,
                                'teams': teams,
                                'control_origination': control_originations,
                                'implementation_status': implementation_status,
                                'parameter': parameter,
                                'responsible_role': responsible_role
                            }
                            create_implementation(new_implementation, certification)


def split_implementations(implementation_details):
    """
    split_implementation(string) -> list of tuples containing (array of teams, solution text)
    Takes a specific control/part's implementation detail section from the SSP, and splits it up by Team. 
    The resulting list is in the following format:
    result = [ ([team a, team b], solution text), ([team c], solution text) ]
    """
    newline_split = implementation_details.split('\n') 
    result = []
    solution_flag = False
    team_list = []
    solution = ''
    for line in newline_split:
        comma_separated = line.split(',')
        if Team.objects.filter(name=comma_separated[0].strip(' ,:')) and ':' in line:
            if solution_flag:
                result.append((team_list, solution))
                team_list = []
                solution = ''
            for team in comma_separated:
                try:
                    team_list.append(Team.objects.get(name__iexact=team.strip(' ,:')))
                except ObjectDoesNotExist:
                    print("Could not find team in database: " + team)


            solution_flag = True
        elif solution_flag:
            solution += line + '\n'
    if team_list:
        result.append((team_list, solution))
    return result

            

        